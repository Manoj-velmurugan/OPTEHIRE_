import docx  # pip install python-docx
import PyPDF2
def read_all_pdf_pages(file_path):
    if file_path.endswith('.pdf'):
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            print(text)
            return text.strip()
    elif file_path.endswith('.txt'):
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read().strip()
    elif file_path.endswith('.docx'):
        doc = docx.Document(file_path)
        print(" ".join([paragraph.text for paragraph in doc.paragraphs]).strip())
        return " ".join([paragraph.text for paragraph in doc.paragraphs]).strip()
    return ""